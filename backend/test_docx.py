import os
from document_service import process_and_store_document, get_all_documents, get_store
import docx

def create_dummy_docx(filename):
    doc = docx.Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It has some text in it.")
    doc.save(filename)

if __name__ == "__main__":
    print("Creating dummy docx...")
    create_dummy_docx("test_dummy.docx")
    
    with open("test_dummy.docx", "rb") as f:
        content = f.read()
        
    print("Processing...")
    try:
        process_and_store_document(content, "test_dummy.docx")
        print("Processed successfully!")
        
        store = get_store()
        print("\nExtracted Text:")
        print(store.get("test_dummy.docx", ""))
        
    except Exception as e:
        print(f"Error: {e}")
        
    # Clean up
    if os.path.exists("test_dummy.docx"):
        os.remove("test_dummy.docx")
